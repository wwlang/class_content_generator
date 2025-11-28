#!/usr/bin/env python3
"""
Markdown to DOCX Converter with Professional Footers.

Converts markdown files to professional DOCX format with:
- Proper heading hierarchy
- Formatted tables
- Lists (bullet and numbered)
- Bold/italic text
- Code blocks
- Block quotes
- Custom footer with course branding and page numbers

Usage:
    python tools/markdown_to_docx.py [course-code] [week-number]
    python tools/markdown_to_docx.py [course-code] syllabus
    python tools/markdown_to_docx.py [course-code] handbook

Example:
    python tools/markdown_to_docx.py BCI2AU 1
    python tools/markdown_to_docx.py BCI2AU syllabus
    python tools/markdown_to_docx.py BCI2AU handbook
"""

import os
import re
import sys
import glob
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# =============================================================================
# CONSTANTS
# =============================================================================

# Template file path (relative to project root)
TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "modern-theme.docx"


# =============================================================================
# DATA STRUCTURES
# =============================================================================

@dataclass
class CourseInfo:
    """Course information for footer branding."""
    code: str
    name: str
    university: str
    campus: str
    instructor: str


@dataclass
class ConversionReport:
    """Result of DOCX conversion."""
    success: bool
    files_converted: int
    files_skipped: int
    output_files: list
    errors: list


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def strip_markdown(text: str) -> str:
    """Remove markdown formatting from text (for footer, etc.)."""
    # Remove bold pairs
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    # Remove italic pairs
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    # Remove inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    # Remove any remaining orphaned asterisks (e.g., leftover from **Label:** Value)
    text = re.sub(r'^\*+\s*', '', text)  # Leading asterisks
    text = re.sub(r'\s*\*+$', '', text)  # Trailing asterisks
    return text.strip()


def ensure_styles_exist(doc: Document) -> None:
    """Ensure required styles exist in the document (for custom templates)."""
    styles = doc.styles

    # Required paragraph styles
    paragraph_styles = [
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
        'List Bullet', 'List Number',
    ]

    for style_name in paragraph_styles:
        try:
            _ = styles[style_name]
        except KeyError:
            styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    # Required table styles
    table_styles = ['Table Grid']

    for style_name in table_styles:
        try:
            _ = styles[style_name]
        except KeyError:
            styles.add_style(style_name, WD_STYLE_TYPE.TABLE)


def set_table_borders(table, top=True, bottom=True, left=True, right=True,
                      inside_h=True, inside_v=True) -> None:
    """
    Set table borders via XML manipulation.

    Args:
        table: The python-docx table object
        top, bottom, left, right: Outer borders (default: True)
        inside_h: Horizontal internal borders (default: True)
        inside_v: Vertical internal borders (default: True)
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing borders element if present
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    # Create new borders element
    tblBorders = OxmlElement('w:tblBorders')

    # Border settings: val="single", sz="4" (half-points), color="auto"
    border_settings = [
        ('w:top', top),
        ('w:left', left),
        ('w:bottom', bottom),
        ('w:right', right),
        ('w:insideH', inside_h),
        ('w:insideV', inside_v),
    ]

    for tag, enabled in border_settings:
        border = OxmlElement(tag)
        if enabled:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # 0.5pt border
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
        else:
            border.set(qn('w:val'), 'nil')
        tblBorders.append(border)

    tblPr.append(tblBorders)


def set_table_cell_margins(table, margin_cm: float = 0.1) -> None:
    """
    Set table cell margins via XML manipulation.

    Args:
        table: The python-docx table object
        margin_cm: Margin in centimeters (default: 0.1 cm)
    """
    # Convert cm to twips (dxa): 1 cm = 567 twips
    margin_twips = int(margin_cm * 567)

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Remove existing cell margin element if present
    existing_margins = tblPr.find(qn('w:tblCellMar'))
    if existing_margins is not None:
        tblPr.remove(existing_margins)

    # Create new cell margins element
    tblCellMar = OxmlElement('w:tblCellMar')

    for side in ['top', 'left', 'bottom', 'right']:
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(margin_twips))
        margin.set(qn('w:type'), 'dxa')
        tblCellMar.append(margin)

    tblPr.append(tblCellMar)


def set_cell_background(cell, rgb_color: tuple) -> None:
    """
    Set cell background color via XML manipulation.

    Args:
        cell: The python-docx table cell object
        rgb_color: RGB color tuple (r, g, b) where each value is 0-255
    """
    # Convert RGB to hex color (without #)
    hex_color = '{:02X}{:02X}{:02X}'.format(*rgb_color)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Remove existing shading element if present
    existing_shd = tcPr.find(qn('w:shd'))
    if existing_shd is not None:
        tcPr.remove(existing_shd)

    # Create new shading element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def ensure_numbering_definitions(doc):
    """
    Ensure the document has proper numbering definitions for bullets and numbered lists.
    This creates the definitions if they don't exist.
    """
    try:
        # Access the numbering part
        if not hasattr(doc, '_part') or not hasattr(doc._part, 'numbering_part'):
            # No numbering part exists, Word will use defaults
            return

        numbering_part = doc._part.numbering_part
        if numbering_part is None:
            return

        # The numbering definitions should already exist in a proper template
        # If issues persist, the template needs to be recreated with proper definitions

    except Exception:
        # If there's any issue accessing numbering, just continue
        # Word will handle it with defaults
        pass


# =============================================================================
# COURSE INFO EXTRACTION
# =============================================================================

def extract_course_info(course_path: str) -> Optional[CourseInfo]:
    """
    Extract course information from course-info.md or syllabus.md.
    """
    # Try course-info.md first
    course_info_path = os.path.join(course_path, "course-info.md")
    syllabus_path = os.path.join(course_path, "syllabus.md")

    content = ""
    if os.path.exists(course_info_path):
        with open(course_info_path, 'r', encoding='utf-8') as f:
            content = f.read()
    elif os.path.exists(syllabus_path):
        with open(syllabus_path, 'r', encoding='utf-8') as f:
            content = f.read()
    else:
        return None

    # Extract course code from folder name
    folder_name = os.path.basename(course_path)
    code_match = re.match(r'^([A-Z0-9]+)', folder_name)
    code = code_match.group(1) if code_match else "COURSE"

    # Extract course name (try table format first, then fallback to colon format)
    name_match = re.search(r'\|\s*\*\*Course Title\*\*\s*\|\s*(.+?)\s*\|', content, re.IGNORECASE)
    if not name_match:
        name_match = re.search(r'(?:Course Title|Title):\s*(.+)', content, re.IGNORECASE)
    name = strip_markdown(name_match.group(1).strip()) if name_match else "Course"

    # Extract university (try table format first, then fallback to colon format)
    uni_match = re.search(r'\|\s*\*\*Program\*\*\s*\|\s*(.+?)\s*\|', content, re.IGNORECASE)
    if not uni_match:
        uni_match = re.search(r'(?:University|Awarding):\s*(.+)', content, re.IGNORECASE)
    university = strip_markdown(uni_match.group(1).strip()) if uni_match else "Andrews University"

    # Extract campus (not in table - look after General Information heading or in header)
    campus_match = re.search(r'–\s*(.+?)\s*Campus', content, re.IGNORECASE)
    if not campus_match:
        campus_match = re.search(r'Campus:\s*(.+)', content, re.IGNORECASE)
    campus = strip_markdown(campus_match.group(1).strip()) if campus_match else "NEU Vietnam"

    # Extract instructor (try table format first, then fallback to colon format)
    instructor_match = re.search(r'\|\s*\*\*Instructor\*\*\s*\|\s*(.+?)\s*\|', content, re.IGNORECASE)
    if not instructor_match:
        instructor_match = re.search(r'Instructor:\s*(.+)', content, re.IGNORECASE)
    instructor = strip_markdown(instructor_match.group(1).strip()) if instructor_match else "Instructor"

    return CourseInfo(
        code=code,
        name=name,
        university=university,
        campus=campus,
        instructor=instructor
    )


# =============================================================================
# FOOTER CREATION
# =============================================================================

def add_header_with_revision_date(doc: Document, source_file_path: str):
    """
    Add header with revision date based on file modification time.

    Format: Revised: [Last Modified Date] (top right)
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = False

    # Get or create header
    header = section.header
    header.is_linked_to_previous = False

    # Clear existing header content
    for para in header.paragraphs:
        para.clear()

    # Get the first paragraph (create if needed)
    if header.paragraphs:
        para = header.paragraphs[0]
    else:
        para = header.add_paragraph()

    # Right align and set spacing
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0

    # Get file modification time
    mod_time = os.path.getmtime(source_file_path)
    mod_date = datetime.fromtimestamp(mod_time).strftime("%B %d, %Y")

    # Add revision date
    run = para.add_run(f"Revised: {mod_date}")
    run.font.size = Pt(9)


def add_footer_with_page_numbers(doc: Document, course_info: CourseInfo, source_file_path: str):
    """
    Add professional footer with course branding, page numbers, and generation date.

    Format:
    Line 1: Course Code Course Name | University | Campus | Instructor | Page X of Y
    Line 2: Generated: [Last Modified Date]
    """
    section = doc.sections[0]

    # Enable different first page if needed
    section.different_first_page_header_footer = False

    # Get or create footer
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing footer content
    for para in footer.paragraphs:
        para.clear()

    # Get the first paragraph (create if needed)
    if footer.paragraphs:
        para = footer.paragraphs[0]
    else:
        para = footer.add_paragraph()

    # Right align and set spacing
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0

    # Add course info text
    run = para.add_run(f"{course_info.code} {course_info.name} | {course_info.university} | {course_info.campus} | {course_info.instructor} | Page ")
    run.font.size = Pt(8)

    # Add PAGE field (current page number)
    add_page_number_field(para)

    # Add " of " text
    run = para.add_run(" of ")
    run.font.size = Pt(8)

    # Add NUMPAGES field (total pages)
    add_num_pages_field(para)

    # Add generation date on second line
    revision_para = footer.add_paragraph()
    revision_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    revision_para.paragraph_format.space_before = Pt(0)
    revision_para.paragraph_format.space_after = Pt(0)
    revision_para.paragraph_format.line_spacing = 1.0

    # Get file modification time
    mod_time = os.path.getmtime(source_file_path)
    mod_date = datetime.fromtimestamp(mod_time).strftime("%B %d, %Y")

    revision_run = revision_para.add_run(f"Generated: {mod_date}")
    revision_run.font.size = Pt(8)


def add_page_number_field(paragraph):
    """Add PAGE field code for current page number."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    run.font.size = Pt(9)


def add_num_pages_field(paragraph):
    """Add NUMPAGES field code for total page count."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

    run.font.size = Pt(9)


# =============================================================================
# MARKDOWN PARSER
# =============================================================================

class MarkdownToDocx:
    """Convert markdown content to DOCX document."""

    def __init__(self, course_info: CourseInfo):
        self.course_info = course_info
        self.doc = None
        self.current_heading = ""  # Track current section heading
        self.in_landscape_section = False  # Track if we're in a landscape section
        self.restart_numbering = False  # Track if we should restart numbering at next list
        # Phase 2: Page break and landscape improvements
        self.current_h2_section = ""  # Track H2 context (Portfolio, Presentation, Rubrics)
        self.is_first_h2 = True  # Skip page break for document title
        self.in_rubrics_section = False  # Different behavior for rubric H3 headings
        self.needs_landscape_for_next_table = False  # Deferred landscape switching

    def convert(self, md_path: str, docx_path: str) -> bool:
        """
        Convert markdown file to DOCX.

        Returns True if successful.
        """
        try:
            # Read markdown content
            with open(md_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # Create new document from template (with fallback to default)
            if TEMPLATE_PATH.exists():
                self.doc = Document(str(TEMPLATE_PATH))
                # Clear paragraphs from template (keep theme/styles/sections)
                while len(self.doc.paragraphs) > 0:
                    p = self.doc.paragraphs[0]._element
                    p.getparent().remove(p)
                # Clear tables from template
                while len(self.doc.tables) > 0:
                    t = self.doc.tables[0]._element
                    t.getparent().remove(t)
                # Ensure required styles exist
                ensure_styles_exist(self.doc)
            else:
                self.doc = Document()
                ensure_styles_exist(self.doc)

            # Set page margins (Top: 2cm, Right: 2cm, Bottom: 2cm, Left: 2.5cm)
            section = self.doc.sections[0]
            section.top_margin = Cm(2)
            section.right_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)

            # Add footer with page numbers and revision date
            add_footer_with_page_numbers(self.doc, self.course_info, md_path)

            # Parse and add content
            self._parse_markdown(content)

            # Save document
            self.doc.save(docx_path)
            return True

        except Exception as e:
            print(f"Error converting {md_path}: {e}")
            return False

    def _parse_markdown(self, content: str):
        """Parse markdown content and add to document."""
        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Check for code block
            if line.strip().startswith('```'):
                i = self._parse_code_block(lines, i)
                continue

            # Check for table
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1] and '---' in lines[i + 1]:
                i = self._parse_table(lines, i)
                continue

            # Check for headers
            if line.startswith('#'):
                self._parse_header(line)
                i += 1
                continue

            # Check for block quote
            if line.strip().startswith('>'):
                i = self._parse_blockquote(lines, i)
                continue

            # Check for bullet list
            if re.match(r'^\s*[-*]\s+', line):
                i = self._parse_bullet_list(lines, i)
                continue

            # Check for numbered list
            if re.match(r'^\s*\d+\.\s+', line):
                i = self._parse_numbered_list(lines, i)
                continue

            # Check for horizontal rule - skip it (heading spacing handles separation)
            if re.match(r'^---+\s*$', line.strip()):
                i += 1
                continue

            # Regular paragraph
            self._add_paragraph(line)
            i += 1

    def _is_assessment_section(self, text: str) -> bool:
        """Check if heading is in an assessment section (Portfolio/Presentation Assessments)."""
        keywords = ['portfolio', 'presentation', 'assessment']
        return any(keyword in text.lower() for keyword in keywords)

    def _parse_header(self, line: str):
        """Parse and add header with smart page breaks and landscape handling."""
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            text_lower = text.lower()

            # ==================== H2 HEADING LOGIC ====================
            if level == 2:
                # Track H2 section context
                self.current_heading = text
                self.current_h2_section = text

                # Check if entering/leaving Rubrics section
                self.in_rubrics_section = 'rubric' in text_lower

                # Return to portrait if leaving rubrics section
                if not self.in_rubrics_section and self.in_landscape_section:
                    section = self.doc.add_section()
                    section.orientation = WD_ORIENT.PORTRAIT
                    section.page_width, section.page_height = section.page_height, section.page_width
                    self.in_landscape_section = False

                # Special handling for Course Calendar in syllabus
                if 'course calendar' in text_lower and not self.in_landscape_section:
                    section = self.doc.add_section()
                    section.orientation = WD_ORIENT.LANDSCAPE
                    section.page_width, section.page_height = section.page_height, section.page_width
                    self.in_landscape_section = True

            # ==================== H3 HEADING LOGIC ====================
            elif level == 3:
                # Restart numbering for ordered lists after H3
                self.restart_numbering = True

                if self.in_rubrics_section:
                    # H3 under "Rubrics" section - set flag for next table to be landscape
                    self.needs_landscape_for_next_table = True
                    # Don't add page break - consecutive rubrics flow together

            # ==================== ADD THE HEADING ====================
            # Map markdown levels to Word heading styles
            style_map = {
                1: 'Heading 1',
                2: 'Heading 2',
                3: 'Heading 3',
                4: 'Heading 4',
                5: 'Heading 5',
                6: 'Heading 6'
            }

            para = self.doc.add_paragraph(style=style_map.get(level, 'Heading 1'))
            self._add_formatted_text(para, text)
            # Override default Word heading spacing
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            # Keep heading with next paragraph
            para.paragraph_format.keep_with_next = True

    def _parse_code_block(self, lines: list, start: int) -> int:
        """Parse code block and return new index."""
        code_lines = []
        i = start + 1

        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1

        # Add code block as monospace paragraph
        code_text = '\n'.join(code_lines)
        para = self.doc.add_paragraph()
        run = para.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        para.paragraph_format.left_indent = Inches(0.5)

        return i + 1  # Skip closing ```

    def _parse_table(self, lines: list, start: int) -> int:
        """Parse markdown table and return new index."""
        # Parse header row
        header_line = lines[start]
        # Split and strip each cell, but keep empty cells to preserve column count
        header_cells = [cell.strip() for cell in header_line.split('|')]
        # Remove ONLY the first and last elements (from | at very start/end of line)
        # Standard markdown tables are: | col1 | col2 | col3 |
        # Which splits to: ['', 'col1', 'col2', 'col3', '']
        if len(header_cells) > 2:  # Need at least ['', content, '']
            header_cells = header_cells[1:-1]  # Remove first and last only

        headers = header_cells  # Keep empty headers

        # Skip separator row
        i = start + 2

        # Parse data rows
        rows = []
        while i < len(lines) and '|' in lines[i] and lines[i].strip():
            row_cells = [cell.strip() for cell in lines[i].split('|')]
            # Remove ONLY the first and last elements (from | at very start/end of line)
            if len(row_cells) > 2:
                row_cells = row_cells[1:-1]

            if row_cells:
                rows.append(row_cells)
            i += 1

        # Check if deferred landscape switch was requested by previous H3 heading
        if self.needs_landscape_for_next_table and not self.in_landscape_section:
            # Add section break and switch to landscape
            section = self.doc.add_section()
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            self.in_landscape_section = True
            # Clear the flag after use
            self.needs_landscape_for_next_table = False

        # Create table (allow empty headers for tables with no header text)
        if len(headers) > 0 and len(rows) > 0:
            table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Table Grid'
            table.allow_autofit = True  # Allow Word to auto-fit columns to content

            # Set table borders: outer borders + horizontal and vertical internal borders
            set_table_borders(table, inside_h=True, inside_v=True)

            # Set table cell margins: 0.1 cm all around
            set_table_cell_margins(table, margin_cm=0.1)

            # Set all cell widths to auto via XML manipulation
            # This makes columns fit their content instead of equal widths
            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:w'), '0')
                    tcW.set(qn('w:type'), 'auto')

            # Add headers (bold, with markdown formatting support and background shading)
            header_cells = table.rows[0].cells
            for j, header in enumerate(headers):
                if j < len(header_cells):
                    # Only apply background shading if header has content
                    if header:
                        set_cell_background(header_cells[j], (226, 232, 240))

                    # Clear default text and use formatted text
                    para = header_cells[j].paragraphs[0]
                    # Set paragraph spacing to 0pt before and after
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.15

                    if header:
                        self._add_formatted_text(para, header)
                        # Make all header text bold and set font size to 8pt
                        for run in para.runs:
                            run.bold = True
                            run.font.size = Pt(8)
                    # else: leave empty cell as is

            # Add data rows (with markdown formatting support and 8pt font)
            for row_idx, row in enumerate(rows):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, cell_text in enumerate(row):
                    if col_idx < len(row_cells):
                        # Use new method that handles <br> tags and bullets
                        self._add_cell_content_with_breaks(row_cells[col_idx], cell_text)

            # Configure table rows
            for idx, row in enumerate(table.rows):
                # Prevent rows from breaking across pages
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                cantSplit = trPr.find(qn('w:cantSplit'))
                if cantSplit is None:
                    cantSplit = OxmlElement('w:cantSplit')
                    trPr.append(cantSplit)

                # Set first row (header) to repeat on each page
                if idx == 0:
                    tblHeader = trPr.find(qn('w:tblHeader'))
                    if tblHeader is None:
                        tblHeader = OxmlElement('w:tblHeader')
                        trPr.append(tblHeader)

        return i

    def _parse_blockquote(self, lines: list, start: int) -> int:
        """Parse blockquote and return new index."""
        quote_lines = []
        i = start

        while i < len(lines) and lines[i].strip().startswith('>'):
            text = lines[i].strip()[1:].strip()  # Remove > prefix
            quote_lines.append(text)
            i += 1

        quote_text = ' '.join(quote_lines)
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.right_indent = Inches(0.5)
        # Set paragraph spacing: 0pt before, 6pt after
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(quote_text)
        run.italic = True
        run.font.size = Pt(10)

        return i

    def _parse_bullet_list(self, lines: list, start: int) -> int:
        """Parse bullet list and return new index."""
        i = start

        while i < len(lines):
            line = lines[i]
            match = re.match(r'^(\s*)[-*]\s+(.+)$', line)

            if not match:
                break

            indent_level = len(match.group(1)) // 2
            text = match.group(2)

            # Check if this is a checkbox item
            checkbox_match = re.match(r'^\[([ xX])\]\s+(.+)$', text)

            para = self.doc.add_paragraph()
            # Set paragraph spacing: 0pt before and after for tight list spacing
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            # Apply manual indentation
            para.paragraph_format.left_indent = Cm(0.5 + (indent_level * 0.5))

            if checkbox_match:
                # This is a checkbox item
                is_checked = checkbox_match.group(1).lower() == 'x'
                checkbox_text = checkbox_match.group(2)

                # Add checkbox symbol as plain text
                checkbox_symbol = '☑' if is_checked else '☐'
                checkbox_run = para.add_run(checkbox_symbol + ' ')
                checkbox_run.font.size = Pt(10)

                # Add the text after the checkbox
                self._add_formatted_text(para, checkbox_text)
            else:
                # Regular bullet item - add bullet as plain text
                bullet_run = para.add_run('• ')
                bullet_run.font.size = Pt(10)

                # Add the text
                self._add_formatted_text(para, text)

            i += 1

        return i

    def _parse_numbered_list(self, lines: list, start: int) -> int:
        """Parse numbered list and return new index."""
        i = start
        list_number = 1  # Track the number ourselves

        # Check if we should restart numbering
        if self.restart_numbering:
            list_number = 1
            self.restart_numbering = False

        while i < len(lines):
            line = lines[i]
            match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)

            if not match:
                break

            indent_level = len(match.group(1)) // 2
            text = match.group(2)

            para = self.doc.add_paragraph()
            # Set paragraph spacing: 0pt before and after for tight list spacing
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15
            # Apply manual indentation
            para.paragraph_format.left_indent = Cm(0.5 + (indent_level * 0.5))

            # Add number as plain text
            number_run = para.add_run(f'{list_number}. ')
            number_run.font.size = Pt(10)

            # Add the text content
            self._add_formatted_text(para, text)

            # Increment number for next item
            list_number += 1

            i += 1

        return i

    def _add_paragraph(self, text: str):
        """Add a regular paragraph with inline formatting."""
        para = self.doc.add_paragraph()
        # Set paragraph spacing: 0pt before, 6pt after
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(6)
        self._add_formatted_text(para, text)

    def _add_formatted_text(self, para, text: str):
        """Add text with inline markdown formatting (bold, italic, code)."""
        # Pattern to match **bold**, *italic*, `code`
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue

            if part.startswith('**') and part.endswith('**'):
                # Bold
                run = para.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(8)
            elif part.startswith('*') and part.endswith('*'):
                # Italic
                run = para.add_run(part[1:-1])
                run.italic = True
                run.font.size = Pt(8)
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = para.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
            else:
                # Regular text
                run = para.add_run(part)
                run.font.size = Pt(8)

    def _add_cell_content_with_breaks(self, cell, cell_text: str):
        """
        Add content to table cell, handling <br> tags and bullet lists.

        Splits on <br> tags and creates separate paragraphs for each segment.
        Lines starting with '- ' are formatted as bullets with small indent.
        """
        # Split on <br> tags (case insensitive)
        segments = re.split(r'<br\s*/?>', cell_text, flags=re.IGNORECASE)

        # Use first paragraph for first segment, create new ones for rest
        for idx, segment in enumerate(segments):
            segment = segment.strip()
            if not segment:
                continue

            # Get or create paragraph
            if idx == 0:
                para = cell.paragraphs[0]
            else:
                para = cell.add_paragraph()

            # Set paragraph spacing to 0pt before and after
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.15

            # Check if this is a bullet item
            if segment.startswith('- '):
                # Remove the bullet marker
                text = segment[2:].strip()

                # Check if this is a checkbox item
                checkbox_match = re.match(r'^\[([ xX])\]\s+(.+)$', text)

                # Add bullet character and text with small indent
                para.paragraph_format.left_indent = Cm(0.3)
                para.paragraph_format.first_line_indent = Cm(-0.2)

                if checkbox_match:
                    # This is a checkbox item
                    is_checked = checkbox_match.group(1).lower() == 'x'
                    checkbox_text = checkbox_match.group(2)

                    # Add checkbox symbol
                    checkbox_symbol = '☑' if is_checked else '☐'
                    checkbox_run = para.add_run(checkbox_symbol + ' ')
                    checkbox_run.font.size = Pt(8)

                    # Add the text after the checkbox
                    self._add_formatted_text(para, checkbox_text)
                else:
                    # Regular bullet
                    bullet_run = para.add_run('• ')
                    bullet_run.font.size = Pt(8)

                    # Add formatted text
                    self._add_formatted_text(para, text)

                # Set font size to 8pt for all runs
                for run in para.runs:
                    run.font.size = Pt(8)
            else:
                # Regular text (may include bold/italic formatting)
                self._add_formatted_text(para, segment)

                # Set font size to 8pt for all runs
                for run in para.runs:
                    run.font.size = Pt(8)


# =============================================================================
# MAIN FUNCTION
# =============================================================================

def convert_week_files(course_code: str, week_number: int, base_path: str = None) -> ConversionReport:
    """
    Convert markdown files in a week folder to DOCX in output/ subfolder.

    Converts:
    - tutorial-content.md -> output/tutorial-content.docx
    - tutorial-tutor-notes.md -> output/tutorial-tutor-notes.docx

    Excludes: lecture-content.md (already becomes slides)
    """
    if base_path is None:
        base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    # Find course folder
    course_folder = os.path.join(base_path, "courses", f"{course_code}*")
    course_matches = glob.glob(course_folder)

    if not course_matches:
        return ConversionReport(
            success=False,
            files_converted=0,
            files_skipped=0,
            output_files=[],
            errors=[f"Course folder not found for: {course_code}"]
        )

    course_path = course_matches[0]
    week_folder = os.path.join(course_path, "weeks", f"week-{week_number:02d}")

    if not os.path.exists(week_folder):
        return ConversionReport(
            success=False,
            files_converted=0,
            files_skipped=0,
            output_files=[],
            errors=[f"Week folder not found: {week_folder}"]
        )

    # Create output folder
    output_folder = os.path.join(week_folder, "output")
    os.makedirs(output_folder, exist_ok=True)

    # Get course info for footer
    course_info = extract_course_info(course_path)
    if not course_info:
        # Use defaults
        course_info = CourseInfo(
            code=course_code,
            name="Course",
            university="Andrews University",
            campus="NEU Vietnam",
            instructor="Instructor"
        )

    converter = MarkdownToDocx(course_info)
    output_files = []
    errors = []
    files_converted = 0
    files_skipped = 0

    # Files to convert (in week folder -> output/)
    week_files = [
        "tutorial-content.md",
        "tutorial-tutor-notes.md"
    ]

    print(f"\nConverting to output/ folder...")
    for md_file in week_files:
        md_path = os.path.join(week_folder, md_file)
        if os.path.exists(md_path):
            docx_file = md_file.replace('.md', '.docx')
            docx_path = os.path.join(output_folder, docx_file)

            print(f"   {md_file} -> output/{docx_file}")
            if converter.convert(md_path, docx_path):
                output_files.append(docx_path)
                files_converted += 1
            else:
                errors.append(f"Failed to convert: {md_file}")
        else:
            print(f"   Skipping (not found): {md_file}")
            files_skipped += 1

    success = files_converted > 0 and not errors

    return ConversionReport(
        success=success,
        files_converted=files_converted,
        files_skipped=files_skipped,
        output_files=output_files,
        errors=errors
    )


def convert_handbook(course_code: str, base_path: str = None) -> ConversionReport:
    """
    Convert assessment-handbook.md to DOCX.
    """
    if base_path is None:
        base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    # Find course folder
    course_folder = os.path.join(base_path, "courses", f"{course_code}*")
    course_matches = glob.glob(course_folder)

    if not course_matches:
        return ConversionReport(
            success=False,
            files_converted=0,
            files_skipped=0,
            output_files=[],
            errors=[f"Course folder not found for: {course_code}"]
        )

    course_path = course_matches[0]
    handbook_md = os.path.join(course_path, "assessment-handbook.md")
    handbook_docx = os.path.join(course_path, "assessment-handbook.docx")

    if not os.path.exists(handbook_md):
        return ConversionReport(
            success=False,
            files_converted=0,
            files_skipped=0,
            output_files=[],
            errors=[f"Assessment handbook not found: {handbook_md}"]
        )

    # Get course info for footer
    course_info = extract_course_info(course_path)
    if not course_info:
        course_info = CourseInfo(
            code=course_code,
            name="Course",
            university="Andrews University",
            campus="NEU Vietnam",
            instructor="Instructor"
        )

    converter = MarkdownToDocx(course_info)

    print(f"\nConverting assessment handbook...")
    print(f"   Input: {handbook_md}")

    if converter.convert(handbook_md, handbook_docx):
        print(f"   Output: {handbook_docx}")
        return ConversionReport(
            success=True,
            files_converted=1,
            files_skipped=0,
            output_files=[handbook_docx],
            errors=[]
        )
    else:
        return ConversionReport(
            success=False,
            files_converted=0,
            files_skipped=0,
            output_files=[],
            errors=["Failed to convert assessment handbook"]
        )


def convert_syllabus(course_code: str, base_path: str = None) -> ConversionReport:
    """
    Convert syllabus.md to DOCX.
    """
    if base_path is None:
        base_path = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    # Find course folder
    course_folder = os.path.join(base_path, "courses", f"{course_code}*")
    course_matches = glob.glob(course_folder)

    if not course_matches:
        return ConversionReport(
            success=False,
            files_converted=0,
            files_skipped=0,
            output_files=[],
            errors=[f"Course folder not found for: {course_code}"]
        )

    course_path = course_matches[0]
    syllabus_md = os.path.join(course_path, "syllabus.md")
    syllabus_docx = os.path.join(course_path, "syllabus.docx")

    if not os.path.exists(syllabus_md):
        return ConversionReport(
            success=False,
            files_converted=0,
            files_skipped=0,
            output_files=[],
            errors=[f"Syllabus not found: {syllabus_md}"]
        )

    # Get course info for footer
    course_info = extract_course_info(course_path)
    if not course_info:
        course_info = CourseInfo(
            code=course_code,
            name="Course",
            university="Andrews University",
            campus="NEU Vietnam",
            instructor="Instructor"
        )

    converter = MarkdownToDocx(course_info)

    print(f"\nConverting syllabus...")
    print(f"   Input: {syllabus_md}")

    if converter.convert(syllabus_md, syllabus_docx):
        print(f"   Output: {syllabus_docx}")
        return ConversionReport(
            success=True,
            files_converted=1,
            files_skipped=0,
            output_files=[syllabus_docx],
            errors=[]
        )
    else:
        return ConversionReport(
            success=False,
            files_converted=0,
            files_skipped=0,
            output_files=[],
            errors=["Failed to convert syllabus"]
        )


def print_report(report: ConversionReport):
    """Print a formatted report."""
    print("\n" + "=" * 60)
    print("DOCX CONVERSION REPORT")
    print("=" * 60)

    status = "SUCCESS" if report.success else "FAILED"
    print(f"\nStatus: {status}")

    print(f"\nSummary:")
    print(f"  - Files converted: {report.files_converted}")
    print(f"  - Files skipped: {report.files_skipped}")

    if report.output_files:
        print(f"\nOutput files:")
        for f in report.output_files:
            print(f"  - {os.path.basename(f)}")

    if report.errors:
        print(f"\nErrors ({len(report.errors)}):")
        for e in report.errors:
            print(f"  - {e}")

    print("\n" + "=" * 60)


# =============================================================================
# CLI ENTRY POINT
# =============================================================================

def main():
    """Command-line entry point."""
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python tools/markdown_to_docx.py [course-code] [week-number]")
        print("  python tools/markdown_to_docx.py [course-code] syllabus")
        print("  python tools/markdown_to_docx.py [course-code] handbook")
        print()
        print("Examples:")
        print("  python tools/markdown_to_docx.py BCI2AU 1")
        print("  python tools/markdown_to_docx.py BCI2AU syllabus")
        print("  python tools/markdown_to_docx.py BCI2AU handbook")
        sys.exit(1)

    course_code = sys.argv[1]
    target = sys.argv[2]

    if target.lower() == 'syllabus':
        print(f"Converting syllabus for {course_code}")
        print("-" * 40)
        report = convert_syllabus(course_code)
    elif target.lower() == 'handbook':
        print(f"Converting assessment handbook for {course_code}")
        print("-" * 40)
        report = convert_handbook(course_code)
    else:
        week_number = int(target)
        print(f"Converting markdown files for {course_code} Week {week_number}")
        print("-" * 40)
        report = convert_week_files(course_code, week_number)

    print_report(report)
    sys.exit(0 if report.success else 1)


if __name__ == "__main__":
    main()
